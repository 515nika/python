import tkinter as tk
from tkinter import messagebox
from particles.electron import specific_charge as electron_charge, compton_wavelength as electron_wavelength
from particles.neutron import specific_charge as neutron_charge, compton_wavelength as neutron_wavelength
from particles.proton import specific_charge as proton_charge, compton_wavelength as proton_wavelength
from docx import Document

def calculate():
    particle = particle_var.get()
    if particle == "Электрон":
        charge = electron_charge()
        wavelength = electron_wavelength()
    elif particle == "Нейтрон":
        charge = neutron_charge()
        wavelength = neutron_wavelength()
    elif particle == "Протон":
        charge = proton_charge()
        wavelength = proton_wavelength()
    else:
        messagebox.showerror("Ошибка", "Выберите частицу")
        return

    result_label.config(text=f"Удельный заряд: {charge} C/kg\nКомптоновская длина волны: {wavelength} m")
    save_to_docx(particle, charge, wavelength)

def save_to_docx(particle, charge, wavelength):
    doc = Document()
    doc.add_heading('Результаты расчёта', 0)
    doc.add_paragraph(f"Частица: {particle}")
    doc.add_paragraph(f"Удельный заряд: {charge} C/kg")
    doc.add_paragraph(f"Комптоновская длина волны: {wavelength} m")
    doc.save(f"{particle}_results.docx")
    messagebox.showinfo("Сохранено", "Результаты сохранены в файл .docx")

root = tk.Tk()
root.title("Расчёт характеристик элементарных частиц")

tk.Label(root, text="Выберите частицу:").pack()
particle_var = tk.StringVar(value="Электрон")
particles = ["Электрон", "Нейтрон", "Протон"]
particle_menu = tk.OptionMenu(root, particle_var, *particles)
particle_menu.pack()

tk.Button(root, text="Рассчитать", command=calculate).pack()

result_label = tk.Label(root, text="", justify=tk.LEFT)
result_label.pack()

root.mainloop()